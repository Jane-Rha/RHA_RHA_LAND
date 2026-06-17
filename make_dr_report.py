"""
DR() Gemini API 토큰 소비량 분석 보고서
- 실측 토큰 반영: 평균 11,536 입력 / 4 출력 (testDRBatch 실측 로그 기준)
- 분석 기간: 2026년 6월 ~ 2027년 6월 (13개월 전망)
- 비용 산출: gemini-3.5-flash 기준 ($1.50/백만 입력, $9.00/백만 출력)
- 오버헤드 ×1.5 포함: 테스트·캐시점검·에러핸들링·개발 목적
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/kevinkim/Desktop/GCX/DR_Token_Analysis.docx"

# ─── 색상 ─────────────────────────────────────────────────────────────────────
def rgb(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))
DARK_H="1F497D"; MID_H="2E75B6"; LGRAY_H="F2F2F2"; WHITE_H="FFFFFF"
DARK=rgb(DARK_H); MID=rgb(MID_H); WHITE=rgb(WHITE_H); BLACK=rgb("000000")
RED=rgb("C00000"); GRAY=rgb("606060"); ORANGE=rgb("833C00")

def shade_cell(cell, h):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    [tcPr.remove(o) for o in tcPr.findall(qn('w:shd'))]
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),h.upper()); tcPr.append(s)

def shade_para(p, h):
    pPr=p._p.get_or_add_pPr()
    [pPr.remove(o) for o in pPr.findall(qn('w:shd'))]
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto')
    s.set(qn('w:fill'),h.upper()); pPr.append(s)

def ap(doc, text, bold=False, sz=11, col=None, italic=False,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6, ind=0):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if ind: p.paragraph_format.left_indent=Cm(ind)
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
    if col: r.font.color.rgb=col
    return p

def heading(doc, text):
    p=ap(doc, text, bold=True, sz=14, col=WHITE, sb=10, sa=4)
    shade_para(p, DARK_H)

def bullet(doc, text, col=None, sz=10):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); r.font.size=Pt(sz)
    if col: r.font.color.rgb=col

def table(doc, hdrs, rows, hbg=DARK_H, alt=LGRAY_H, bold_last=False):
    t=doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(hdrs):
        c=t.rows[0].cells[i]; c.text=h; shade_cell(c, hbg)
        r=c.paragraphs[0].runs[0]; r.bold=True; r.font.size=Pt(9); r.font.color.rgb=WHITE
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    for ri,row in enumerate(rows):
        last=bold_last and ri==len(rows)-1
        bg=DARK_H if last else (alt if ri%2==0 else WHITE_H)
        fg=WHITE if last else BLACK
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val); shade_cell(c, bg)
            r=c.paragraphs[0].runs[0]; r.font.size=Pt(9); r.font.color.rgb=fg; r.bold=last
            c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# 핵심 상수 (실측 기반)
# ═══════════════════════════════════════════════════════════════════════════════
INP_PRICE = 1.50    # $/백만 입력 토큰 (gemini-3.5-flash)
OUT_PRICE = 9.00    # $/백만 출력 토큰 (gemini-3.5-flash)

# testDRBatch 실측: 19회 호출, 입력 219,189 tok → 평균 11,536 tok/회
IN_TOK    = 11_536  # 실측 평균 입력 토큰
OUT_TOK   = 4       # 실측 평균 출력 토큰 (68 / 19 ≈ 3.6 → 4)

OVERHEAD  = 1.5     # 테스트·캐시점검·에러핸들링·개발 목적 오버헤드
BYPASS    = 0.10    # 키워드 패스트패스 우회율 (~10%)

CPC = (IN_TOK * INP_PRICE + OUT_TOK * OUT_PRICE) / 1_000_000  # 호출당 기본 비용
cpr = CPC * (1 - BYPASS)                                        # 리뷰당 비용 (우회 제외)

# ── 월간 속도 (실측/추정) ─────────────────────────────────────────────────────
S26_PM   = 362   # Galaxy S26: 실측 1,267행 / 3.5개월
YJH_PM   = 297   # 유지훈P: 실측 1,635행 / 5.5개월
I17E_PM  = 13    # iPhone 17e: GDrive 40행 / 3개월
Z8_PM    = 180   # Galaxy Z8: 예상 (Z7 유사)
IPH18_PM = 350   # iPhone 18: 예상 (S26 유사)
S27_PM   = 362   # Galaxy S27: 예상 (S26 유사)
IPH18E_PM= 13    # iPhone 18e: 예상 (17e 유사)
REG_PM   = 60    # 정규 시트 합계 (SDA 31 + Auto_Acc 16 + Power_Acc 10 + 전략폰 3)

# ── 2026.06 ~ 2027.06 월별 데이터 ────────────────────────────────────────────
# 컬럼: (레이블, S26, YJH, 17e, Z8, iPh18, S27, iPh18e, REG)
months = [
    ('2026년 6월',  S26_PM, YJH_PM, I17E_PM,  0,          0,          0,       0,         REG_PM),
    ('2026년 7월',  0,      YJH_PM, 0,         Z8_PM,      0,          0,       0,         REG_PM),
    ('2026년 8월',  0,      YJH_PM, 0,         Z8_PM,      0,          0,       0,         REG_PM),
    ('2026년 9월',  0,      YJH_PM, 0,         Z8_PM,      IPH18_PM//2,0,       0,         REG_PM),
    ('2026년 10월', 0,      YJH_PM, 0,         0,          IPH18_PM,   0,       0,         REG_PM),
    ('2026년 11월', 0,      YJH_PM, 0,         0,          IPH18_PM,   0,       0,         REG_PM),
    ('2026년 12월', 0,      YJH_PM, 0,         0,          IPH18_PM//2,0,       0,         REG_PM),
    ('2027년 1월',  0,      YJH_PM, 0,         0,          0,          0,       0,         REG_PM),
    ('2027년 2월',  0,      YJH_PM, 0,         0,          0,          S27_PM,  0,         REG_PM),
    ('2027년 3월',  0,      YJH_PM, 0,         0,          0,          S27_PM,  0,         REG_PM),
    ('2027년 4월',  0,      YJH_PM, 0,         0,          0,          S27_PM,  IPH18E_PM//2, REG_PM),
    ('2027년 5월',  0,      YJH_PM, 0,         0,          0,          0,       IPH18E_PM, REG_PM),
    ('2027년 6월',  0,      YJH_PM, 0,         0,          0,          0,       IPH18E_PM, REG_PM),
]

# ═══════════════════════════════════════════════════════════════════════════════
# 문서 생성
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin=sec.bottom_margin=Cm(1.8)
    sec.left_margin=sec.right_margin=Cm(2.0)

# ── 제목 ──────────────────────────────────────────────────────────────────────
ap(doc, 'Gemini API 토큰 소비량 분석 보고서',
   bold=True, sz=18, col=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
ap(doc, 'DR() 커스텀 함수 — Spigen GCX CS 리뷰 모니터링 시스템',
   sz=12, col=MID, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
ap(doc, '작성일: 2026년 6월 15일  |  분석 기간: 2026년 6월 ~ 2027년 6월 (13개월 전망)',
   sz=9, col=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
for note in [
    '※ 비용 산출: gemini-3.5-flash 기준 ($1.50/백만 입력토큰, $9.00/백만 출력토큰) — 상한선 추정',
    '※ 입력 토큰: testDRBatch 실측 평균 11,536 tok/회 (19회 호출, 219,189 tok 합산)',
    '※ 오버헤드 ×1.5 적용: 테스트·캐시 점검·에러 핸들링·개발 목적 포함',
]:
    ap(doc, note, sz=10, col=RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
doc.add_paragraph()

# ══ §1 실측 토큰 분석 ═════════════════════════════════════════════════════════
heading(doc, '1. testDRBatch 실측 토큰 분석')
ap(doc,
   '2026년 6월 Galaxy S26 1-3점 시트에서 testDRBatch 함수를 실행한 결과. '
   '모델: gemini-3.5-flash (STATUS 200 확인). 총 19회 Gemini 호출, 1회 캐시 히트.',
   sz=10, sa=4)

# 실측 데이터 발췌
sample = [
    ('Row 3',  '11,518', '4', '11,522', '자력약함',   'magnet is pretty weak…'),
    ('Row 4',  '11,579', '2', '11,581', '형합',       'いつもはspigen のタフアーマーを…'),
    ('Row 5',  '11,557', '2', '11,559', '형합',       'the case does not fit the phone…'),
    ('Row 7',  '11,511', '4', '11,515', '자력약함',   'very weak magnets and because of…'),
    ('Row 10', '11,616', '3', '11,619', '카툭튀',     'really wanted to like this case…'),
    ('Row 15', '11,513', '5', '11,518', '부착어려움', 'my samsung s26 ultra screen protector…'),
    ('Row 18', '11,505', '1', '11,506', '가격',       '품질 좋은 슈피겐 제품이지만 가격이…'),
    ('Row 21', '11,590', '5', '11,595', '부착어려움', '한국회사 한국 제품인데 설명서는 왜…'),
]
table(doc,
    ['행', '입력 토큰', '출력 토큰', '합계', '분류 결과', '리뷰 본문 (앞 50자)'],
    [list(r) for r in sample])
doc.add_paragraph()
table(doc,
    ['구분', '수치', '비고'],
    [
        ['총 Gemini 호출', '19회', '20행 중 1회 캐시 히트'],
        ['총 입력 토큰',   '219,189 tok', 'testDRBatch 합산'],
        ['총 출력 토큰',   '68 tok', ''],
        ['호출당 평균 입력', f'{219189//19:,} tok', '219,189 ÷ 19 = 11,536'],
        ['호출당 평균 출력', f'{68//19} tok', '68 ÷ 19 ≈ 3.6 → 4'],
        ['기존 추정 (문자 수 기반)', '8,604 tok', '실측 대비 -34% 과소 추정'],
        ['실측 vs 추정 차이', '+34%', '한국어 토크나이저 특성, 프롬프트 오버헤드 등'],
    ], bold_last=False)

# ══ §2 요금 및 호출당 비용 ════════════════════════════════════════════════════
heading(doc, '2. Gemini API 요금 및 호출당 비용 (실측 기반)')
ap(doc,
   '모든 비용은 gemini-3.5-flash 요금으로 보수적 상한선 산정. '
   '현재 실제 운영 모델(gemini-3.1-flash-lite)과는 요금이 다릅니다.',
   sz=10, sa=4)
table(doc,
    ['모델', '구분', '$/백만 입력', '$/백만 출력',
     '호출당 기본 비용', '×1.5 후 비용', '리뷰당 비용(×1.5,우회10%)'],
    [
        ['gemini-3.5-flash', '비용 산출 기준 (상한)',
         '$1.50', '$9.00', f'${CPC:.6f}', f'${CPC*OVERHEAD:.6f}', f'${cpr*OVERHEAD:.6f}'],
        ['gemini-3.1-flash-lite', '현재 실 운영 모델 (참고)',
         '$0.25', '$1.50',
         f'${(IN_TOK*0.25+OUT_TOK*1.50)/1e6:.6f}',
         f'${(IN_TOK*0.25+OUT_TOK*1.50)/1e6*OVERHEAD:.6f}', '—'],
        ['gemini-2.5-flash-lite', '현재 폴백 (참고)',
         '$0.10', '$0.40',
         f'${(IN_TOK*0.10+OUT_TOK*0.40)/1e6:.6f}',
         f'${(IN_TOK*0.10+OUT_TOK*0.40)/1e6*OVERHEAD:.6f}', '—'],
    ])
doc.add_paragraph()
ap(doc,
   f'gemini-3.5-flash 호출 1회 기본: ${CPC:.6f}  →  ×1.5 오버헤드 적용 후: ${CPC*OVERHEAD:.6f}\n'
   f'(입력 {IN_TOK:,}tok × $1.50/백만 = ${IN_TOK*INP_PRICE/1e6:.6f}  +  '
   f'출력 {OUT_TOK}tok × $9.00/백만 = ${OUT_TOK*OUT_PRICE/1e6:.6f})',
   sz=10, col=RED, bold=True, sa=4)

# ══ §3 시트별 월간 속도 기준 ════════════════════════════════════════════════════
heading(doc, '3. 시트별 월간 속도 기준')
ap(doc,
   '주력 기종(Galaxy S26, 유지훈P)은 실제 스프레드시트 확인치 기반. '
   '미출시 기종(Z8, iPhone 18, S27, 18e)은 유사 기종 실측치 기반 추정.',
   sz=10, sa=4)
table(doc,
    ['시트 / 기종', '유형', '월간 속도', '근거', '모니터링 기간'],
    [
        ['Galaxy S26',  '주력 기종 (종료 예정)', f'{S26_PM}행/월',  '★ 실측: 1,267행 / 3.5개월', '2026년 3~6월'],
        ['유지훈P',      '일상 운영 (연중)',      f'{YJH_PM}행/월',  '★ 실측: 1,635행 / 5.5개월', '2026년 전체'],
        ['iPhone 17e',  '주력 기종 (종료 예정)', f'{I17E_PM}행/월', 'GDrive: 40행 / 3개월',       '2026년 3~6월'],
        ['Galaxy Z8',   '주력 기종 (예정)',       f'{Z8_PM}행/월',   '추정: Galaxy Z7 유사 규모',   '2026년 7~9월'],
        ['iPhone 18',   '주력 기종 (예정)',       f'{IPH18_PM}행/월','추정: S26 유사 규모',         '2026년 9월~12월'],
        ['Galaxy S27',  '주력 기종 (예정)',       f'{S27_PM}행/월',  '추정: S26 확인치와 동일',     '2027년 2~4월'],
        ['iPhone 18e',  '주력 기종 (예정)',       f'{IPH18E_PM}행/월','추정: 17e 확인치와 동일',   '2027년 4~6월'],
        ['SDA',         '정규 (배포 예정)',        '~31행/월',        'GDrive 날짜 범위 역산',       '연중'],
        ['Auto_Acc',    '정규 (배포 예정)',        '~16행/월',        'GDrive 날짜 범위 역산',       '연중'],
        ['Power_Acc',   '정규 (배포 예정)',        '~10행/월',        'GDrive 73행 / 7개월',         '연중'],
        ['전략폰',       '정규 (배포 예정)',        '~3행/월',         'GDrive 36행 / 19개월',        '연중'],
        ['정규 합계',    '일상 운영',              f'{REG_PM}행/월',  'SDA+Auto_Acc+Power_Acc+전략폰', '연중'],
    ], bold_last=True)

# ══ §4 월별 타임라인 — 2026.06 ~ 2027.06 ═══════════════════════════════════════
heading(doc, '4. 월별 DR() 호출 타임라인 (2026년 6월 ~ 2027년 6월)')
ap(doc,
   '★ = 실측 확인치. † = 예정/추정. 반월(½) 출시는 월 속도의 절반 적용. '
   '비용 열은 gemini-3.5-flash 기준 기본 비용 (오버헤드 ×1.5 미포함).',
   sz=10, sa=4)

h_mon = ['월', 'S26★', 'YJH★', 'iPh17e', 'Z8†', 'iPh18†', 'S27†', 'iPh18e†',
         '정규', '합계', '입력 토큰(K)', '기본 비용']
rows_mon = []
grand_calls = 0
grand_cost  = 0.0
grand_tok   = 0

for r in months:
    lbl, s26, yjh, i17e, z8, i18, s27, i18e, reg = r
    tot   = s26+yjh+i17e+z8+i18+s27+i18e+reg
    cost  = tot * CPC
    intok = tot * IN_TOK
    grand_calls += tot; grand_cost += cost; grand_tok += intok
    def f(v): return str(v) if v else '—'
    rows_mon.append([lbl, f(s26), f(yjh), f(i17e), f(z8), f(i18), f(s27), f(i18e), f(reg),
                     str(tot), f"{intok/1000:.0f}K", f"${cost:.3f}"])

rows_mon.append(['13개월 합계',
                 f"~{sum(r[1] for r in months):,}",
                 f"~{sum(r[2] for r in months):,}",
                 f"~{sum(r[3] for r in months)}",
                 f"~{sum(r[4] for r in months)}",
                 f"~{sum(r[5] for r in months):,}",
                 f"~{sum(r[6] for r in months):,}",
                 f"~{sum(r[7] for r in months)}",
                 f"~{sum(r[8] for r in months):,}",
                 f"{grand_calls:,}",
                 f"{grand_tok/1000:.0f}K",
                 f"${grand_cost:.2f}"])
table(doc, h_mon, rows_mon, bold_last=True)
doc.add_paragraph()
bullet(doc, '★ S26: 2026년 6월 모니터링 종료 예정 (4개월 합산 기준)')
bullet(doc, '★ 유지훈P: 연중 운영 지속 (13개월 × 297행/월)')
bullet(doc, '† Galaxy Z8: 2026년 7월 출시 예상, 3개월 모니터링 (Jul~Sep), 약 180행/월')
bullet(doc, '† iPhone 18: 2026년 9월 중순 출시 예상 → 9월 반월(175) + 10~11월 풀(350) + 12월 반월(175)')
bullet(doc, '† Galaxy S27: 2027년 2월 출시 예상, 3개월 모니터링 (Feb~Apr), S26과 동일 속도 적용')
bullet(doc, '† iPhone 18e: 2027년 4월 출시 예상, 3개월 모니터링 (Apr 반월~Jun), 17e와 동일 속도 적용')
bullet(doc, '정규 시트 4종 합계 60행/월은 현재 미배포 — 배포 시 해당 월부터 합산')

# ══ §5 비용 요약 및 오버헤드 ══════════════════════════════════════════════════
heading(doc, '5. 비용 요약 — 기본 vs 오버헤드 ×1.5')
grand_oh = grand_cost * OVERHEAD
grand_real_model = grand_oh / 6   # 3.1-flash-lite 기준 (3.5 대비 약 1/6)

table(doc,
    ['구분', 'DR() 호출 수', '입력 토큰 합계', '기본 비용 (USD)', '×1.5 후 최종 비용 (USD)'],
    [
        ['2026년 6월 ~ 12월 (7개월)',
         f"{sum(sum(r[1:]) for r in months[:7]):,}",
         f"{sum(sum(r[1:]) for r in months[:7])*IN_TOK:,}",
         f"${sum(sum(r[1:]) for r in months[:7])*CPC:.2f}",
         f"${sum(sum(r[1:]) for r in months[:7])*CPC*OVERHEAD:.2f}"],
        ['2027년 1월 ~ 6월 (6개월)',
         f"{sum(sum(r[1:]) for r in months[7:]):,}",
         f"{sum(sum(r[1:]) for r in months[7:])*IN_TOK:,}",
         f"${sum(sum(r[1:]) for r in months[7:])*CPC:.2f}",
         f"${sum(sum(r[1:]) for r in months[7:])*CPC*OVERHEAD:.2f}"],
        ['2026.06 ~ 2027.06 전체 (13개월)',
         f"{grand_calls:,}",
         f"{grand_tok:,}",
         f"${grand_cost:.2f}",
         f"${grand_oh:.2f}"],
    ], bold_last=True)
doc.add_paragraph()
ap(doc,
   f'최종 예상 비용 (오버헤드 ×1.5 포함): USD ${grand_oh:.2f}',
   bold=True, sz=14, col=DARK, sa=2, align=WD_ALIGN_PARAGRAPH.CENTER)
ap(doc,
   f'기본 ${grand_cost:.2f}  ×  1.5  =  ${grand_oh:.2f}  |  총 {grand_calls:,}건 호출  |  {grand_tok/1e6:.2f}백만 입력 토큰',
   sz=11, col=GRAY, sa=4, align=WD_ALIGN_PARAGRAPH.CENTER)

# 오버헤드 항목별 분해
doc.add_paragraph()
ap(doc, '오버헤드 구성 (×1.5 = 기본 100% + 추가 50%)', bold=True, sz=11, sa=2)
table(doc,
    ['오버헤드 항목', '내용', '비중 (추정)'],
    [
        ['testDRBatch 실행', 'GAS 에디터에서 배치 테스트 실행, 캐시 미적중 호출 발생', '~20%'],
        ['캐시 무효화 후 재확인', 'DR_CACHE_VERSION 변경 시 전체 캐시 초기화 → 전량 재호출', '~15%'],
        ['에러 핸들링 재시도', '1차 모델 실패 시 2차 모델 폴백 호출 (각 행 토큰 2배)', '~10%'],
        ['DEBUG_DR 수동 호출', '개발 디버그, 분류 결과 검증 목적 수동 실행', '~5%'],
        ['합계 추가 오버헤드', '기본 비용의 약 50% 추가 = ×1.5 배율', '×1.5'],
    ], bold_last=True)

# ══ §6 기존 추정 vs 실측 비교 ═════════════════════════════════════════════════
heading(doc, '6. 기존 토큰 추정 vs 실측 비교')
ap(doc,
   '기존 분석은 문자 수 / 3 방식으로 토큰을 추정했으나, '
   'Gemini API의 실제 토크나이저(SentencePiece BPE)는 한국어·다국어 혼합 시 '
   '문자 수 대비 더 많은 토큰을 할당합니다.',
   sz=10, sa=4)
old_cpc = (8604 * 1.50 + 4 * 9.00) / 1_000_000
old_total_calls = 10822  # 이전 보고서 2025~2026 합계
table(doc,
    ['항목', '기존 추정', '실측 기반 (이번 보고서)', '차이'],
    [
        ['입력 토큰/회', '8,604 tok', f'{IN_TOK:,} tok', '+2,932 tok (+34%)'],
        ['출력 토큰/회', '4 tok', '4 tok', '변동 없음'],
        ['호출당 기본 비용', f'${old_cpc:.6f}', f'${CPC:.6f}', f'+${CPC-old_cpc:.6f} (+34%)'],
        ['호출당 ×1.5 비용', f'${old_cpc*OVERHEAD:.6f}', f'${CPC*OVERHEAD:.6f}', f'+${(CPC-old_cpc)*OVERHEAD:.6f}'],
        ['추정 방법', '문자 수 ÷ 3 (추산)', 'testDRBatch 19회 실측 평균', '실측 우선'],
    ])

# ══ §7 주요 가정 및 유의사항 ═══════════════════════════════════════════════════
heading(doc, '7. 주요 가정 및 유의사항')
for t, b in [
    ('비용 산출 기준',
     f'gemini-3.5-flash ($1.50/백만 입력, $9.00/백만 출력) 기준 보수적 상한선. '
     f'현재 실 운영 모델(gemini-3.1-flash-lite)로는 약 ${grand_real_model:.2f} 예상 (약 83% 절감).'),
    ('실측 토큰',
     f'testDRBatch 19회 실측: 평균 {IN_TOK:,} 입력 / 4 출력. '
     '기존 문자 수 기반 추정(8,604)보다 34% 높음. '
     '향후 Defect 시트 항목 수 변경 시 토큰도 비례 변동.'),
    ('오버헤드 ×1.5',
     '테스트 실행, 캐시 초기화 후 재처리, 에러 재시도, 개발 목적 수동 호출 포함. '
     '1.5배는 보수적 추정이며 개발 주기·배포 빈도에 따라 달라질 수 있음.'),
    ('키워드 패스트패스',
     '약 10% 리뷰가 heavy/bulky/yellow/button/attach/scratch 키워드 우회 → 0토큰. '
     '월간 비용 = 호출 수 × $0.017340 × 0.90 (우회 10% 제외) × 1.5 (오버헤드).'),
    ('6시간 캐시',
     '동일 리뷰 본문은 6시간 내 재계산 시 캐시 적중 → 0토큰. '
     'DR_CACHE_VERSION 변경(현재 v22) 시 전체 캐시 초기화 → 전량 재호출.'),
    ('미출시 기종 속도 추정',
     'Z8·iPh18·S27·iPh18e의 월간 속도는 유사 기종 실측치(S26·17e)를 기준으로 추정. '
     '실제 판매량·리뷰 수에 따라 ±30% 오차 가능.'),
    ('정규 시트 배포 미반영',
     'SDA, Auto_Acc, Power_Acc, 전략폰은 현재 DR() 미사용. '
     '배포 시 월 +60건 추가 (×1.5 오버헤드 포함 시 월 약 +$5.50 추가).'),
    ('주말 미운영',
     'Master.gs dailyJob()은 평일만 실행 (약 22일/월). '
     '월간 속도 추정은 달력일 기준이 아닌 실제 수집 건수 기반.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
    r1=p.add_run(f'{t}: '); r1.bold=True; r1.font.size=Pt(10)
    r2=p.add_run(b); r2.font.size=Pt(10)

# ══ §8 총괄 요약 ══════════════════════════════════════════════════════════════
heading(doc, '8. 총괄 요약')
table(doc,
    ['구분', 'DR() 호출', '입력 토큰', '기본 비용', '×1.5 최종 비용',
     '(참고) 현재 모델 기준'],
    [
        ['2026년 6월~12월 (7개월)',
         f"{sum(sum(r[1:]) for r in months[:7]):,}",
         f"{sum(sum(r[1:]) for r in months[:7])*IN_TOK/1e6:.2f}M",
         f"${sum(sum(r[1:]) for r in months[:7])*CPC:.2f}",
         f"${sum(sum(r[1:]) for r in months[:7])*CPC*OVERHEAD:.2f}",
         f"${sum(sum(r[1:]) for r in months[:7])*CPC*OVERHEAD/6:.2f}"],
        ['2027년 1월~6월 (6개월)',
         f"{sum(sum(r[1:]) for r in months[7:]):,}",
         f"{sum(sum(r[1:]) for r in months[7:])*IN_TOK/1e6:.2f}M",
         f"${sum(sum(r[1:]) for r in months[7:])*CPC:.2f}",
         f"${sum(sum(r[1:]) for r in months[7:])*CPC*OVERHEAD:.2f}",
         f"${sum(sum(r[1:]) for r in months[7:])*CPC*OVERHEAD/6:.2f}"],
        ['전체 (2026.06 ~ 2027.06)',
         f"{grand_calls:,}",
         f"{grand_tok/1e6:.2f}M",
         f"${grand_cost:.2f}",
         f"${grand_oh:.2f}",
         f"${grand_real_model:.2f}"],
    ], bold_last=True)

doc.add_paragraph()
ap(doc,
   f'gemini-3.5-flash 기준  |  오버헤드 ×1.5 포함  |  최종 예상 비용: USD ${grand_oh:.2f}',
   bold=True, sz=14, col=DARK, sa=2, align=WD_ALIGN_PARAGRAPH.CENTER)
ap(doc,
   f'총 DR() 호출: {grand_calls:,}건  |  총 입력 토큰: {grand_tok/1e6:.2f}M개  |  '
   f'월 평균: {grand_calls/13:.0f}건/월, ${grand_oh/13:.2f}/월',
   sz=11, col=GRAY, sa=4, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
ap(doc,
   f'• 호출 1회 기본: ${CPC:.6f}  →  ×1.5 후: ${CPC*OVERHEAD:.6f}\n'
   f'• 리뷰 1건당 (키워드 10% 우회 + ×1.5): ${cpr*OVERHEAD:.6f}\n'
   f'• 유지훈P + 정규 시트 상시 운영 월간 비용 (×1.5): '
   f'${(YJH_PM+REG_PM)*CPC*OVERHEAD:.4f}/월\n'
   f'• 참고 — 현재 실제 모델(gemini-3.1-flash-lite) 기준 오버헤드 포함: '
   f'약 ${grand_real_model:.2f} (약 83% 절감)',
   sz=10, sa=4)
ap(doc,
   '※ Google AI Studio pay-as-you-go 요금 기준. 프리 티어 및 약정 할인 미적용. '
   '실제 청구 금액 ±30% 내외 변동 가능.',
   sz=9, col=GRAY)

doc.add_paragraph()
fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run(
    'Generated by Claude Code (claude-sonnet-4-6) | Spigen GCX 자동화팀 | 2026년 6월 15일')
fr.font.size=Pt(8); fr.font.color.rgb=rgb("999999"); fr.italic=True

doc.save(OUTPUT)
print(f"저장 완료 → {OUTPUT}")
print(f"\n[ 핵심 수치 ]")
print(f"  실측 토큰/회: 입력 {IN_TOK:,}, 출력 {OUT_TOK}")
print(f"  호출당 기본 비용 (3.5-flash): ${CPC:.6f}")
print(f"  호출당 ×1.5 후:              ${CPC*OVERHEAD:.6f}")
print(f"  13개월 총 호출:              {grand_calls:,}건")
print(f"  기본 비용 합계:              ${grand_cost:.2f}")
print(f"  ×1.5 오버헤드 최종:          ${grand_oh:.2f}")
print(f"  (현재 모델 기준 참고):        ${grand_real_model:.2f}")
